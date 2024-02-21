from xlsx2html import xlsx2html
from bs4 import BeautifulSoup
import pprint
import html
import os
import conf
import argparse
import re

import openpyxl 
import http.server
import socketserver 
import webbrowser
from docx2python import docx2python



'''''''''''''''''''''''''''''''''
    Utility functions here
'''''''''''''''''''''''''''''''''

# is the cell valid?
def is_valid_cell(s):
    if '"$"' in s.string:
        s.string=s.string.replace('"',"")
    if '($' in s.string:
        s.string= s.string.replace('(',"-").replace(')',"")
    return s is not None and s.string[0] in conf.valid_chars

# get new cell value from old cell value
def process_cell(td, sheet_name, name, context):
    id =td['id']
    col = id[len(sheet_name) + 1]
    
    # Process context
    if name in conf.d_to_i_contexts and context and context[0] == 'D':
        context = "I" + context[1:]
    try:
        outstring = td.string
        outstring_int = outstring
        # Process $
        dollar = "" if outstring[0] != "$" else "$"
        outstring = outstring[1:] if dollar else outstring
        outstring_int = outstring_int[1:] if dollar else outstring_int

        # Process -
        if outstring[0] == "-":
            if len(outstring) > 1 and outstring[1] == '$':
                outstring_int = outstring_int[1:]
                minus = "-$"
            else:
                minus = "-"
            sign = ' sign="-"'
            outstring = outstring_int[1:]

        else:
            minus = ""
            sign = ""

        if td.string in ["-", "$ -"]:
            minus = ""
            value = ''
            format = 'ixt:fixed-zero'
            outstring = "-"
        else:
            format = 'ixt:num-dot-decimal'
            value = ""

        row = id[(len(sheet_name)+2):]
        cell_id = col + row
        id = f"{sheet_name.replace(' ', '_')}_{cell_id}"
        
        content = f'{minus}{dollar}<ix:nonFraction contextRef="{context}" name="{name}" unitRef="USD" id="{id}" decimals="0" format="{format}"{sign}{value}>' + \
                  f'{outstring}</ix:nonFraction>'
        return content
    except Exception as e:
        return None
   
'''''''''''''''''''''''''''''''''
    Main code here
'''''''''''''''''''''''''''''''''

# Parse command line first
parser = argparse.ArgumentParser()
parser.add_argument('--i', type=str, metavar="input_file(xlsx)", required=True, help="Input xlsx file name")
parser.add_argument('--o', type=str, metavar="output_file(html)", help="Output html file name")

args = parser.parse_args()

if args.i:
    input_file = args.i
if args.o:
    output_file = args.o
else:
    output_file = input_file.split(".")[0] + ".html"

# Parse Context first
try:
    xlsx2html("contexts.xlsx", "temp")
except:
    print("No context file!")
    exit(0)
with open("temp", 'r') as f:
    html = f.read() 

# Remove header
html = html.replace(conf.original_header, '')
# Analyze
soup = BeautifulSoup(html, 'html.parser')
context_skip_lines = 1
i = 0
context_name_map = {}
context_ref_map = {}
for tr in soup.find_all('tr'):
    # Skip some lines
    if i < context_skip_lines:
        i += 1
        continue

    tds = tr.findAll('td')
    if len(tds) == 5:
        scope = tds[0].string.strip().lower()
        statement = tds[1].string.strip().lower()
        header = tds[2].string.strip().lower()
        name = tds[3].string.strip()
        ref = pprint.pformat(tds[4].contents)
        ref_index = ref.find("<xbrli:context") 
       
        if ref_index == -1: #critical error : invalid xbrli content
            continue
        ref = ref[ref_index:-1] 

        # calculate index
        index = f"{scope}@{statement}" 
        if statement != "statement of activities": # General type context sheet
            # Add to context name map
            if not index in context_name_map:
                context_name_map[index] = {}
            context_name_map[index][header] = name
            # Add to context ref map
            if not index in context_ref_map:
                context_ref_map[index] = {}
            context_ref_map[index][header] = ref
        else: # Special type context sheet
            if not index in context_ref_map:
                context_ref_map[index] = {}
            context_ref_map[index][name] = ref
            

sheet_count = 0
html_in = ""
ix_header_content = ""
ix_header_names = []
html=''

for ws in openpyxl.load_workbook(input_file).sheetnames:
    xlsx2html(input_file,'temp0',sheet=ws)
    with open('temp0', "r") as f:
        html += f.read()

# Get rid of header generated by the xlsx2html library
html_trunc = html.replace(conf.original_header, '')
soup = BeautifulSoup(html_trunc, 'html.parser')
sheet_name = None
name = None # name property of context, eg:acfr:Reve...
statement_scope = None
header_row = None # row of header of the table
context_name = None # context name: for special sheets
special_name_map = {} # "column" ->"context name" map
special_sheet = False # special sheet eg Statement of Activities
special_rows = [] # green rows

for td in soup.find_all('td'):
    id = td['id']
    if sheet_name is None or sheet_name != id.split("!")[0]:
        sheet_name = id.split("!")[0]   
        header_cols = []   # columns of headers 
        header_titles = {} # dict of header titles "col" -> "title"
        new_sheet = 0
        if sheet_name == "Statement of Activities":
            special_sheet = True 
        else: 
            special_sheet = False 
  
    if is_valid_cell(td):
        td['style'] = td['style'] + ';text-align:right; font-size:12.5px'
    else:
        td['style'] = td['style'] + ';font-size:14px'

    # if is_valid_cell(td):
    #     td['style'] = td['style'] + ';text-align:right; font-size:12.5px'
    # else:
    #     td['style'] = td['style'] + ';'

    col = id[len(sheet_name) + 1]
    row = int(id[(len(sheet_name) + 2):]) 
    if row == header_row:
        header_titles[col] = td.string.strip().lower() 
    
    if  header_row and new_sheet==0 and row > header_row:
        if special_sheet:
            if not statement_scope in context_ref_map:
                print(f"Invalid statement scope {statement_scope} on sheet {sheet_name} ", statement_scope)
            for key, value in context_ref_map[statement_scope].items():
                if not key in ix_header_names:
                    ix_header_content += f"\n{value}" 
                    ix_header_names.append(key)
        else:
            for header,ref in header_titles.items():
                if header and  ref and statement_scope and "xbrl element" not in ref:
                    try:
                        ix_name = context_name_map[statement_scope][ref]
                        if not ix_name in ix_header_names: # avoid duplicate context names
                            ix_header_content += f"\n{context_ref_map[statement_scope][ref]}"  
                            ix_header_names.append(ix_name)
                    except:
                        print(f"Invalid ix header for {statement_scope} on sheet {sheet_name}")
        new_sheet=1
    if col + str(row) == "B2":
        statement_scope = td.string.strip().lower()
    if col + str(row)  == "B3":
        statement_scope = statement_scope + "@" + td.string.strip().lower()
    
    content = td.string.strip()
    if content == "XBRL Element":
        if col == 'A':
            header_row = row
        header_cols.append(col)
    if header_row and row > header_row: 
        if col in header_cols: 
            name = content 
    
    if special_sheet: 
        if re.match("[DI][0-9]{8}[A-Za-z_]*", content):#context name
            special_name_map[col] = content
            if not row in special_rows: # Remember green rows
                special_rows.append(row)
        if name and is_valid_cell(td) and col in special_name_map:
            if not header_row:
                print("No XBRL Element header")  
                exit(0)
            content = process_cell(td, sheet_name, name, special_name_map[col])
            if content:
                td.string = content
    else:
        if col > 'B' and name and is_valid_cell(td):
            try:
                ht = header_titles[col].strip().lower()
                context_name = context_name_map[statement_scope][ht]
            except:
                context_name = 'I20220630'
                print(f"Invalid scope for {id}, {statement_scope}, {ht}")
            content = process_cell(td, sheet_name, name, context_name)
            if content:
                td.string = content
    new_id = id.replace(" ", "").replace("!", "_")
    td['id'] = new_id
    
    if header_cols == [] and len(td.get_text(strip=True)) == 0:
        background = td['style'].split(';')
        if 'background-color' in background[0]: 
            td.decompose()

    if (col in header_cols) or (special_sheet and row in special_rows): 
        td.decompose() 

word_file_path = 'Xbrl Independent Auditor.docx'
def extract_text_from_docx(file_path):
    doc_html_result = docx2python(file_path, html = True)
    text_content = ""
    for page in doc_html_result.body[0][0][0]:
        for paragraph in page:
            for element in paragraph:
                if not isinstance(element, dict) or ('type' in element and element['type'] != 'footer'):
                    if isinstance(element, str):
                        text_content += element       
        text_content += '\n'     
    doc_html_result.close()
    return text_content

word_text = extract_text_from_docx(word_file_path)
# print(word_text)
html_in += soup.prettify("utf-8").decode("utf-8")

word_text = re.sub(r'style="[^"]*"', '', word_text)

def add_paragraph_tags(html_input):
    soup = BeautifulSoup(html_input, 'html.parser')
    for line in soup.find_all(string=True):
        if line.strip() and line.parent.name not in ['h3','h4','span','b','u','i']:
            paragraphs = line.strip().split('\n') # Split text into paragraphs based on double newlines
            for paragraph in paragraphs:
                p_tag = soup.new_tag("p")
                p_tag.string = paragraph.strip()
                line.insert_before(p_tag) # Insert the <p> tag before the current line
                line.insert_before("\n")  # Add a newline after the <p> tag for formatting 
            line.extract()
    return str(soup).replace('</p><i>','<i>').replace('</p>\n<i>','\n<i>').replace('</i><p>','</i>')
    

word_text = add_paragraph_tags(word_text)
from docx import Document
def extract_images(docx_file):
    doc = Document(docx_file)
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            image_name = rel.target_ref[:]  # Get the image name
            images.append((image_name, image_data))
    return images

# Example usage
images = extract_images(word_file_path)
for idx, (image_name, image_data) in enumerate(images):
    with open(f"{image_name}", "wb") as f:
        f.write(image_data)

def replace_image_placeholders(content):
    pattern = r'----(.*?)----'
    replaced_content = re.sub(pattern, r'<img src="\1" style="width:300px;" />', content)
    return replaced_content

word_text = replace_image_placeholders(word_text)

ix_header = conf.ix_header_start + ix_header_content + conf.ix_header_end  +  '<ix:nonNumeric> \n <h2 align="center"> AUDITOR\'S REPORT </h2>  '+ word_text+'</ix:nonNumeric>'
html_out = conf.new_header + '\n' + ix_header.replace("$place_id$", conf.place_id) + '\n<ix:nonNumeric> \n <h2 align="center"> FINANCIAL REPORT </h2> </ix:nonNumeric>'

for line in html_in.splitlines(): 
    html_out = html_out + line + '\n'

html_out = html_out  +'</body></html>' 
html_out = html_out.replace("&lt;", "<")
html_out = html_out.replace("&gt;", ">")

html_out = html_out.replace("xbrli:startdate", "xbrli:startDate")
html_out = html_out.replace("xbrli:enddate", "xbrli:endDate")

os.remove("temp0")
os.remove("temp")
encoded_html_out = html_out.encode('utf-8')
decoded_html_out = encoded_html_out.decode('utf-8') # Decode the encoded string using UTF-8 encoding

with open(output_file, 'w', encoding='utf-8') as f:
    f.write(decoded_html_out)
    print(f"Successfully converted to {output_file}")


os.system('"C:\\Users\\Mitrah154\\Downloads\\arelle-win\\arelleCmdLine" -- file=D:\\Process-XBRL-main\\ca_clayton_2022.html --plugins EdgarRenderer')
os.system(f'"C:\\Users\\Mitrah154\\Downloads\\arelle-win\\arelleCmdLine" --plugins=ixbrl-viewer -f {output_file} --save-viewer {output_file}')

handler = http.server.SimpleHTTPRequestHandler
httpd = socketserver.TCPServer(("", 8081), handler)

webbrowser.open(f"http://localhost:8081/{output_file}")
httpd.serve_forever()
